from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Pt, RGBColor

from bs4 import NavigableString, Tag


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run_element = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run_element.append(r_pr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.append(text_element)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def _append_runs(paragraph, node, bold=False, italic=False, underline=False) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline
        return

    if not isinstance(node, Tag):
        return

    tag = node.name.lower()
    child_bold = bold or tag in {"strong", "b"}
    child_italic = italic or tag in {"em", "i"}
    child_underline = underline or tag in {"u"}

    if tag == "a":
        href = node.get("href")
        link_text = node.get_text("", strip=True) or href or ""
        if href:
            _add_hyperlink(paragraph, href, link_text)
        elif link_text:
            paragraph.add_run(link_text)
        return

    if tag == "br":
        paragraph.add_run("\n")
        return

    for child in node.children:
        _append_runs(paragraph, child, child_bold, child_italic, child_underline)


def append_html_to_document(document, root: Tag) -> None:
    block_tags = {
        "p",
        "div",
        "section",
        "article",
        "blockquote",
        "pre",
        "h1",
        "h2",
        "h3",
        "h4",
        "h5",
        "h6",
    }
    list_tags = {"ul", "ol"}

    for child in root.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                document.add_paragraph(text)
            continue

        if not isinstance(child, Tag):
            continue

        tag = child.name.lower()
        if tag in list_tags:
            _append_list(document, child)
            continue

        if tag in block_tags or child.get_text(strip=True):
            paragraph = document.add_paragraph()
            if tag.startswith("h") and len(tag) == 2 and tag[1].isdigit():
                level = int(tag[1])
                for run in paragraph.runs:
                    run.bold = True
                paragraph.style = f"Heading {min(level, 9)}"
            _append_runs(paragraph, child)


def _append_list(document, list_node: Tag) -> None:
    ordered = list_node.name.lower() == "ol"
    for index, item in enumerate(list_node.find_all("li", recursive=False), start=1):
        prefix = f"{index}. " if ordered else "• "
        paragraph = document.add_paragraph()
        paragraph.add_run(prefix)
        for child in item.children:
            _append_runs(paragraph, child)


def add_separator_paragraph(document, char: str = "-", count: int = 49) -> None:
    paragraph = document.add_paragraph(char * count)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(160, 160, 160)
        run.font.size = Pt(9)
