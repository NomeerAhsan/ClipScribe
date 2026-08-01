import os
import tempfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor

from backend.formatter import add_separator_paragraph, append_html_to_document
from backend.parser import parse_html_fragment
from backend.utils import duplicate_key, load_state, save_state


class DocumentWriter:
    def __init__(self, document_path: Path, state_path: Path) -> None:
        self.document_path = document_path
        self.state_path = state_path
        self.document_path.parent.mkdir(parents=True, exist_ok=True)

    def _load_document(self) -> Document:
        if self.document_path.exists():
            return Document(str(self.document_path))
        document = Document()
        title = document.add_heading("ClipScribe Research Notes", level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph("")
        return document

    def _format_timestamp(self, captured_at: str) -> str:
        try:
            normalized = captured_at.replace("Z", "+00:00")
            dt = datetime.fromisoformat(normalized)
            return dt.strftime("%d %B %Y · %I:%M %p")
        except ValueError:
            return captured_at

    def _article_exists(self, document: Document, page_url: str) -> bool:
        marker = f"Source: {page_url}"
        for paragraph in document.paragraphs:
            if marker in paragraph.text:
                return True
        return False

    def _append_article_header(
        self,
        document: Document,
        page_title: str,
        page_url: str,
        captured_at: str,
    ) -> None:
        add_separator_paragraph(document, "=", 49)
        document.add_heading(page_title, level=1)
        source = document.add_paragraph()
        source.add_run("Source: ").bold = True
        source.add_run(page_url)
        captured = document.add_paragraph()
        captured.add_run("Captured: ").bold = True
        captured.add_run(self._format_timestamp(captured_at))
        add_separator_paragraph(document, "=", 49)
        document.add_paragraph("")

    def append_highlight(
        self,
        html: str,
        page_title: str,
        page_url: str,
        captured_at: str,
    ) -> tuple[str, int]:
        key = duplicate_key(page_url, html)
        state = load_state(self.state_path)

        if key in state["duplicate_keys"]:
            return "duplicate", 0

        document = self._load_document()
        if not self._article_exists(document, page_url):
            self._append_article_header(document, page_title, page_url, captured_at)

        article_counts = state["article_counts"]
        highlight_number = article_counts.get(page_url, 0) + 1

        label = document.add_paragraph()
        label_run = label.add_run(f"Highlight {highlight_number}")
        label_run.bold = True
        label_run.font.size = Pt(11)

        root = parse_html_fragment(html).select_one(".clipscribe-root")
        if root is not None:
            append_html_to_document(document, root)
        else:
            document.add_paragraph(html)

        add_separator_paragraph(document)

        self._save_document(document)

        state["duplicate_keys"].append(key)
        state["article_counts"][page_url] = highlight_number
        save_state(self.state_path, state)

        return "saved", highlight_number

    def _save_document(self, document: Document) -> None:
        self.document_path.parent.mkdir(parents=True, exist_ok=True)

        fd, temp_path = tempfile.mkstemp(suffix=".docx", dir=self.document_path.parent)
        os.close(fd)
        temp_file = Path(temp_path)

        try:
            document.save(str(temp_file))
            os.replace(temp_file, self.document_path)
        except PermissionError as exc:
            if temp_file.exists():
                temp_file.unlink(missing_ok=True)
            raise PermissionError(
                "Could not save Research.docx. Close the file in Microsoft Word "
                "(or reload it after saving), then try highlighting again."
            ) from exc
        except OSError as exc:
            if temp_file.exists():
                temp_file.unlink(missing_ok=True)
            raise OSError(
                "Could not write to Research.docx. Close Microsoft Word if it has "
                "the file open, then try again."
            ) from exc
